import os
import uuid

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

from app.config import settings


class DocumentService:

    def __init__(self):

        self.output_folder = os.path.join(
            settings.MEDIA_FOLDER,
            "generated"
        )

        os.makedirs(
            self.output_folder,
            exist_ok=True
        )

    def generate_docx(
        self,
        title: str,
        content: str
    ):

        filename = f"{uuid.uuid4()}.docx"

        filepath = os.path.join(
            self.output_folder,
            filename
        )

        doc = Document()

        doc.add_heading(title, level=1)

        doc.add_paragraph(content)

        doc.save(filepath)

        return filepath

    def generate_pdf(
        self,
        title: str,
        content: str
    ):

        filename = f"{uuid.uuid4()}.pdf"

        filepath = os.path.join(
            self.output_folder,
            filename
        )

        styles = getSampleStyleSheet()

        pdf = SimpleDocTemplate(filepath)

        story = [
            Paragraph(f"<b>{title}</b>", styles["Heading1"]),
            Paragraph(content.replace("\n", "<br/>"), styles["BodyText"])
        ]

        pdf.build(story)

        return filepath


document_service = DocumentService()